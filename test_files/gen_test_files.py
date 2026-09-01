# -*- coding: utf-8 -*-
"""生成智读·文件配置测试文件（6 种格式，内容为可映射配置草稿的套餐方案文档）。

参照 docs/testdata/智读测试方案_家庭融合.md：
每份文件含 3 条套餐（A 合规通过 / B 零资费待修正 / C 缺月费待修正），
可经 rd_file_parse → batchFromDocument 映射出 3 条配置草稿，支持多文件批量导入。
"""
import csv
import os

from docx import Document
import pandas as pd
from PIL import Image, ImageDraw, ImageFont

OUT = os.path.dirname(os.path.abspath(__file__))

# 各格式一致的套餐正文（正则切分锚点：套餐X：…；要素用「；」分隔）
PLAN_MD = """# 智慧社区融合套餐方案（智读测试稿）

> 用途：产商品研发助手 · 智读·文件配置 联调
> 用法：上传本文件后发送「帮我导入这份智慧社区融合方案」

## 一、方案背景

面向家庭宽带用户推出融合套餐组合，覆盖主套餐与加装包，一次文档映射出多套配置草稿并完成合规校验。

## 二、套餐清单

套餐A：智慧社区融合畅享158；月费158元；含40GB流量+500分钟语音+500M宽带；目标家庭；全渠道销售；有合约12个月；不可重复订购

套餐B：智慧社区体验0元流量包；月费0元；含5GB流量；目标家庭；全渠道；无合约；可重复订购；折扣100%

套餐C：智慧社区融合加装包；目标家庭；电渠+厅店；依赖宽带；未写月费

## 三、预期结果

| 套餐 | 预期状态 | 说明 |
|------|----------|------|
| A 智慧社区融合畅享158 | 合规通过 | 字段完整，可作为入库样例 |
| B 智慧社区体验0元流量包 | 待修正 | 零资费 + 无合约 + 可重复，触发高风险规则 |
| C 智慧社区融合加装包 | 待修正 | 缺月费、依赖未填全 |
"""

PLAN_TXT = """帮我导入这份智慧社区融合方案：

套餐A：智慧社区融合畅享158；月费158元；含40GB流量+500分钟语音+500M宽带；目标家庭；全渠道销售；有合约12个月；不可重复订购

套餐B：智慧社区体验0元流量包；月费0元；含5GB流量；目标家庭；全渠道；无合约；可重复订购；折扣100%

套餐C：智慧社区融合加装包；目标家庭；电渠+厅店；依赖宽带；未写月费
"""

PLAN_ROWS = [
    ("A", "智慧社区融合畅享158", "158元", "40GB流量+500分钟语音+500M宽带", "家庭", "全渠道销售", "有合约12个月", "不可重复订购"),
    ("B", "智慧社区体验0元流量包", "0元", "5GB流量", "家庭", "全渠道", "无合约", "可重复订购"),
    ("C", "智慧社区融合加装包", "未写月费", "依赖宽带", "家庭", "电渠+厅店", "-", "-"),
]

PLAN_DOCX = (
    "智慧社区融合套餐方案（智读测试稿）\n"
    "\n"
    "一、方案背景\n"
    "面向家庭宽带用户推出融合套餐组合，覆盖主套餐与加装包，要求一次文档映射出多套配置草稿并完成合规校验。\n"
    "\n"
    "二、套餐清单\n"
    "套餐A：智慧社区融合畅享158；月费158元；含40GB流量+500分钟语音+500M宽带；目标家庭；全渠道销售；有合约12个月；不可重复订购\n"
    "套餐B：智慧社区体验0元流量包；月费0元；含5GB流量；目标家庭；全渠道；无合约；可重复订购；折扣100%\n"
    "套餐C：智慧社区融合加装包；目标家庭；电渠+厅店；依赖宽带；未写月费\n"
)

PLAN_PDF = [
    "智慧社区融合套餐方案（智读测试稿）",
    "",
    "一、方案背景",
    "面向家庭宽带用户推出融合套餐组合，覆盖主套餐与加装包，一次文档映射出多套配置草稿并完成合规校验。",
    "",
    "二、套餐清单",
    "套餐A：智慧社区融合畅享158；月费158元；含40GB流量+500分钟语音+500M宽带；目标家庭；全渠道销售；有合约12个月；不可重复订购",
    "套餐B：智慧社区体验0元流量包；月费0元；含5GB流量；目标家庭；全渠道；无合约；可重复订购；折扣100%",
    "套餐C：智慧社区融合加装包；目标家庭；电渠+厅店；依赖宽带；未写月费",
]


def gen_md():
    with open(os.path.join(OUT, "智慧社区融合方案.md"), "w", encoding="utf-8") as f:
        f.write(PLAN_MD)


def gen_txt():
    with open(os.path.join(OUT, "智慧社区融合方案_粘贴版.txt"), "w", encoding="utf-8") as f:
        f.write(PLAN_TXT)


def gen_docx():
    doc = Document()
    doc.add_heading("智慧社区融合套餐方案（智读测试稿）", level=0)
    doc.add_heading("一、方案背景", level=1)
    doc.add_paragraph(
        "面向家庭宽带用户推出融合套餐组合，覆盖主套餐与加装包，"
        "要求一次文档映射出多套配置草稿并完成合规校验。"
    )
    doc.add_heading("二、套餐清单", level=1)
    table = doc.add_table(rows=1, cols=8)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["序号", "资费名称", "月费", "包含资源", "目标客群", "销售渠道", "合约", "重复性"]):
        hdr[i].text = h
    for row in PLAN_ROWS:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    doc.add_heading("三、套餐原文", level=1)
    doc.add_paragraph(PLAN_DOCX.split("二、套餐清单\n")[1].split("## 三")[0].strip())
    doc.save(os.path.join(OUT, "智慧社区融合方案.docx"))


def gen_pdf():
    """PIL 绘制文本页 + img2pdf 转换，保证中文正常显示"""
    font_path = r"C:\Windows\Fonts\msyh.ttc"
    try:
        font = ImageFont.truetype(font_path, 22)
        title_font = ImageFont.truetype(font_path, 30)
    except OSError:
        font = ImageFont.load_default()
        title_font = font
    img = Image.new("RGB", (1400, 900), "white")
    draw = ImageDraw.Draw(img)
    y = 40
    for i, line in enumerate(PLAN_PDF):
        draw.text((60, y), line, fill="black", font=title_font if i == 0 else font)
        y += 50 if i == 0 else 42
    img_path = os.path.join(OUT, "_tmp_pdf_page.png")
    img.save(img_path)

    import img2pdf
    with open(os.path.join(OUT, "智慧社区融合方案.pdf"), "wb") as f:
        f.write(img2pdf.convert(img_path))
    os.remove(img_path)


def gen_xlsx():
    """每行一条套餐，一行即一个套餐段落（TSV 输出后仍可按行切分）"""
    df = pd.DataFrame(PLAN_ROWS, columns=[
        "序号", "资费名称", "月费", "包含资源", "目标客群", "销售渠道", "合约", "重复性",
    ])
    path = os.path.join(OUT, "智慧社区融合方案.xlsx")
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="套餐清单", index=False)


def gen_csv():
    path = os.path.join(OUT, "智慧社区融合方案.csv")
    with open(path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["序号", "资费名称", "月费", "包含资源", "目标客群", "销售渠道", "合约", "重复性"])
        for row in PLAN_ROWS:
            writer.writerow(row)


if __name__ == "__main__":
    for name in ("产品说明文档.md", "使用说明文本.txt", "测试Word文档.docx",
                 "测试PDF文档.pdf", "测试Excel表格.xlsx", "测试CSV数据.csv"):
        p = os.path.join(OUT, name)
        if os.path.exists(p):
            os.remove(p)
    gen_md()
    gen_txt()
    gen_docx()
    gen_pdf()
    gen_xlsx()
    gen_csv()
    for name in sorted(os.listdir(OUT)):
        p = os.path.join(OUT, name)
        if os.path.isfile(p) and name != os.path.basename(__file__):
            print(f"{name}  {os.path.getsize(p)} bytes")
