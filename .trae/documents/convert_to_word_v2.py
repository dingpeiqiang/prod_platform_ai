"""
将专利申请书 Markdown 文件转换为 Word (.docx) 格式 v2
逐行状态机方式处理，确保内容不丢失
"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_text_run(paragraph, text, bold=False, font_size=12, font_name='宋体', color=None, italic=False):
    """向段落添加文本片段"""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = color
    return run


def parse_inline_formatting(text):
    """解析行内格式化标记，返回 (parts) 列表 """
    # 先处理超链接 [text](url) -> text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    
    # 分割加粗文本 **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    result = []
    for p in parts:
        if p.startswith('**') and p.endswith('**'):
            result.append(('bold', p[2:-2]))
        else:
            result.append(('normal', p))
    return result


def add_formatted_paragraph(doc, text, font_size=12, font_name='宋体', 
                             alignment=None, indent=0, space_before=3, space_after=3,
                             line_spacing=22, color=None):
    """添加一个支持行内格式的段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(line_spacing)
    if indent > 0:
        p.paragraph_format.left_indent = Cm(indent)
    if alignment:
        p.alignment = alignment
    
    formatted_parts = parse_inline_formatting(text)
    for fmt, txt in formatted_parts:
        is_bold = (fmt == 'bold')
        add_text_run(p, txt, bold=is_bold, font_size=font_size, font_name=font_name, color=color)
    return p


def convert_md_to_word_v2(md_file, docx_file):
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # 状态变量
    in_code_block = False
    code_block_lines = []
    code_block_lang = ''
    
    in_table = False
    table_lines = []
    
    in_ordered_list = False
    
    i = 0
    total = len(lines)
    
    while i < total:
        raw_line = lines[i]
        line = raw_line.rstrip('\n\r')  # 保留行尾空格用于表格识别
        stripped = line.strip()
        
        # --- 空行 ---
        if not stripped:
            # 如果在代码块中，保留空行
            if in_code_block:
                code_block_lines.append('')
            # 如果在表格中，表格结束
            elif in_table:
                process_table_v2(doc, table_lines)
                table_lines = []
                in_table = False
            # 如果在有序列表中，结束列表
            elif in_ordered_list:
                in_ordered_list = False
            i += 1
            continue
        
        # --- 代码块处理 (```) ---
        if stripped.startswith('```'):
            if not in_code_block:
                # 开始代码块，记录语言
                in_code_block = True
                code_block_lang = stripped[3:].strip()
                code_block_lines = []
            else:
                # 结束代码块，写入
                in_code_block = False
                if code_block_lines:
                    # 添加代码块
                    for code_line in code_block_lines:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.left_indent = Cm(0.3)
                        add_text_run(p, code_line, font_size=8.5, font_name='Consolas', 
                                      color=RGBColor(0x33, 0x33, 0x33))
                    # 代码块后空行
                    doc.add_paragraph()
                code_block_lines = []
                code_block_lang = ''
            i += 1
            continue
        
        # 如果正在代码块中，直接收集
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # --- 水平线 ---
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            # 添加一个分隔段落（细水平线效果用段落间距模拟）
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # 添加一条细线
            run = p.add_run('─' * 50)
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            i += 1
            continue
        
        # --- 表格 ---
        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            table_lines.append(stripped)
            i += 1
            # 继续收集下一行
            continue
        elif in_table:
            # 表格已结束，处理表格
            if table_lines:
                process_table_v2(doc, table_lines)
                table_lines = []
            in_table = False
            # 不i++，让当前行继续被后续处理
        
        # --- 标题 ---
        if stripped.startswith('### '):
            text = stripped[4:]
            add_formatted_paragraph(doc, text, font_size=14, font_name='黑体',
                                     space_before=10, space_after=5)
            i += 1
            continue
        
        if stripped.startswith('## '):
            text = stripped[3:]
            add_formatted_paragraph(doc, text, font_size=15, font_name='黑体',
                                     space_before=14, space_after=6)
            i += 1
            continue
        
        if stripped.startswith('# '):
            text = stripped[2:]
            add_formatted_paragraph(doc, text, font_size=18, font_name='黑体',
                                     alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                     space_before=18, space_after=10)
            i += 1
            continue
        
        # --- 无序列表 ---
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            # 检查是否是 **加粗** 开头
            formatted_parts = parse_inline_formatting(text)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.8)
            # 添加圆点符号
            run = p.add_run('• ')
            run.font.size = Pt(12)
            
            for fmt, txt in formatted_parts:
                is_bold = (fmt == 'bold')
                add_text_run(p, txt, bold=is_bold, font_size=12)
            i += 1
            continue
        
        # --- 有序列表 (1. xxx) ---
        match = re.match(r'^(\d+)[\.\)]\s+(.*)', stripped)
        if match:
            num = match.group(1)
            text = match.group(2)
            formatted_parts = parse_inline_formatting(text)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.8)
            # 添加编号
            run = p.add_run(f'{num}. ')
            run.font.size = Pt(12)
            
            for fmt, txt in formatted_parts:
                is_bold = (fmt == 'bold')
                add_text_run(p, txt, bold=is_bold, font_size=12)
            
            in_ordered_list = True
            i += 1
            continue
        
        # --- (a) (b) (c) 格式列表 ---
        match = re.match(r'^\(([a-zA-Z])\)\s+(.*)', stripped)
        if match:
            letter = match.group(1)
            text = match.group(2)
            formatted_parts = parse_inline_formatting(text)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.8)
            run = p.add_run(f'({letter}) ')
            run.font.size = Pt(12)
            
            for fmt, txt in formatted_parts:
                is_bold = (fmt == 'bold')
                add_text_run(p, txt, bold=is_bold, font_size=12)
            i += 1
            continue
        
        # --- 普通段落 ---
        # 检查是否以 **加粗开头** 如 **1. xxx** 或 **标题**
        formatted_parts = parse_inline_formatting(stripped)
        
        # 如果整个段落只有加粗部分，视为小标题
        all_bold = all(fmt == 'bold' for fmt, _ in formatted_parts)
        
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = Pt(22)
        
        if all_bold:
            # 加粗段落视为小标题
            for fmt, txt in formatted_parts:
                add_text_run(p, txt, bold=True, font_size=12, font_name='黑体')
        else:
            for fmt, txt in formatted_parts:
                is_bold = (fmt == 'bold')
                add_text_run(p, txt, bold=is_bold, font_size=12)
        
        i += 1
    
    # 处理未结束的表格
    if in_table and table_lines:
        process_table_v2(doc, table_lines)
    
    # 保存
    doc.save(docx_file)
    print(f"✅ Word 文档已生成: {docx_file}")


def process_table_v2(doc, table_lines):
    """处理Markdown表格"""
    # 过滤掉分隔行 (|---|---|)
    data_lines = [l for l in table_lines if not re.match(r'^\|[\s\-:]+\|', l)]
    
    if len(data_lines) < 2:
        return
    
    # 解析表头
    headers = []
    for cell in data_lines[0].split('|')[1:-1]:
        headers.append(cell.strip())
    
    # 解析数据行
    rows = []
    for line in data_lines[1:]:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
    
    if not rows:
        return
    
    # 确保列数一致
    max_cols = max(len(headers), max(len(r) for r in rows))
    while len(headers) < max_cols:
        headers.append('')
    
    # 创建表格
    table = doc.add_table(rows=len(rows) + 1, cols=max_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        add_text_run(p, header, bold=True, font_size=10, font_name='黑体')
        set_cell_shading(cell, "D9E2F3")
    
    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < max_cols:
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                add_text_run(p, cell_text, font_size=10)
    
    doc.add_paragraph()  # 表格后空行


if __name__ == '__main__':
    md_file = r'D:\工作\sitech\项目\研发\git_workspace\AI\prod_platform_ai\.trae\documents\patent_01_dynamic_form_llm_inference.md'
    docx_file = r'D:\工作\sitech\项目\研发\git_workspace\AI\prod_platform_ai\.trae\documents\patent_01_dynamic_form_llm_inference.docx'
    convert_md_to_word_v2(md_file, docx_file)